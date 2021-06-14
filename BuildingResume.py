from docx import Document
from docx.shared import Inches
import pyttsx3

#Speak Method
def speak(name):
    pyttsx3.speak(name)


document = Document()

document.add_picture('pic.jpg', width=Inches(2.0))

name = input("What is your name: ")
speak('Hello' + name + 'How are you today?')
phoneNo = input("Phone Number? ")
Address = input("Address: ")
Email = input("Email? ")
document.add_paragraph(
    name + ' | '+ phoneNo + ' | '+ Address 
)

#about me
document.add_heading('About me')
document.add_paragraph(input('Tell about yourself? '))


#work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter Company')
from_date = input('From Date')
to_date = input('To Date')

p.add_run(company + ' ').bold = True
p.add_run(from_date+ '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company)
p.add_run(experience_details)

#more experiences
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company')
        from_date = input('From Date')
        to_date = input('To Date')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date+ '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company)
        p.add_run(experience_details )
    else:
        break

#List of Skills
skills = input('What are your skills? ' )
p.add_run('-' + ' ' + skills + '\n')
while True:
    has_more_skills = input('Do you have more skills? yes or no: ')
    if has_more_skills.lower() == 'yes':
        skills = input('What are your skills? ' )
        p.add_run('-' + ' ' + skills + '\n')
    else:
        break

#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using mojola's code"


document.save('resume.docx') 